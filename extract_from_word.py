#!/usr/bin/env python3
"""
Extract content from Word documents and convert to JSON format for the website.

This script reads .docx files from a 'word-docs' directory and converts them
to the JSON format used by content.json, or directly updates HTML files.
"""

import os
import json
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches


def extract_text_from_docx(docx_path):
    """Extract structured content from a Word document."""
    doc = Document(docx_path)
    
    content = {
        "title": "",
        "intro": "",
        "sections": []
    }
    
    current_section = None
    paragraphs = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        # Check if paragraph is a heading based on style
        style_name = para.style.name.lower()
        is_heading = (
            'heading' in style_name or
            'title' in style_name or
            para.style.name.startswith('Heading')
        )
        
        # Also check if text looks like a heading (short, bold, etc.)
        if not is_heading:
            # Check if it's formatted as bold and short (likely a heading)
            runs = para.runs
            if runs and all(run.bold for run in runs if run.text.strip()):
                if len(text) < 100:  # Short text that's all bold
                    is_heading = True
        
        if is_heading:
            # Save previous section if exists
            if current_section:
                if paragraphs:
                    current_section["body"] = "\n\n".join(paragraphs)
                content["sections"].append(current_section)
            
            # Start new section
            current_section = {"heading": text}
            paragraphs = []
            
            # Check heading level
            if 'heading 1' in style_name or 'title' in style_name:
                if not content["title"]:
                    content["title"] = text
            elif 'heading 2' in style_name:
                # This might be the main title
                if not content["title"]:
                    content["title"] = text
        else:
            # Regular paragraph
            if not content["title"] and not current_section:
                # First paragraph might be title or intro
                if len(text) < 200:
                    if not content["title"]:
                        content["title"] = text
                    elif not content["intro"]:
                        content["intro"] = text
                else:
                    content["intro"] = text
            elif not current_section:
                # Content before first section
                if not content["intro"]:
                    content["intro"] = text
                else:
                    content["intro"] += "\n\n" + text
            else:
                paragraphs.append(text)
    
    # Save last section
    if current_section:
        if paragraphs:
            current_section["body"] = "\n\n".join(paragraphs)
        content["sections"].append(current_section)
    
    return content


def extract_images_from_docx(docx_path, output_dir):
    """Extract images from Word document and save them."""
    doc = Document(docx_path)
    images = []
    
    # Get the document's image parts
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            try:
                image_part = rel.target_part
                image_ext = image_part.content_type.split('/')[-1]
                if image_ext not in ['png', 'jpeg', 'jpg', 'gif']:
                    continue
                
                image_filename = f"extracted_{len(images)}_{rel.rId}.{image_ext}"
                image_path = os.path.join(output_dir, image_filename)
                
                with open(image_path, 'wb') as f:
                    f.write(image_part.blob)
                
                images.append({
                    "filename": image_filename,
                    "path": image_path,
                    "rel_id": rel.rId
                })
            except Exception as e:
                print(f"Warning: Could not extract image {rel.rId}: {e}")
    
    return images


def map_filename_to_content_key(filename):
    """Map Word document filename to content.json key."""
    filename_lower = filename.lower()
    
    mapping = {
        'maps': 'maps',
        'location': 'maps',
        'cartographer': 'maps',
        'tectonic': 'tectonics',
        'earthquake': 'tectonics',
        'volcano': 'tectonics',
        'weathering': 'weathering',
        'erosion': 'weathering',
        'mass': 'weathering',
        'fluvial': 'fluvial',
        'ocean': 'fluvial',
        'coastline': 'fluvial',
        'climate': 'climate',
        'biome': 'climate',
        'about': 'about',
        'home': 'home',
        'reference': 'references',
    }
    
    for key, value in mapping.items():
        if key in filename_lower:
            return value
    
    return None


def update_content_json(content_json_path, new_content, content_key):
    """Update content.json with new content."""
    if os.path.exists(content_json_path):
        with open(content_json_path, 'r', encoding='utf-8') as f:
            content_data = json.load(f)
    else:
        content_data = {}
    
    # Merge new content with existing
    if content_key in content_data:
        # Preserve existing structure, update with new content
        existing = content_data[content_key]
        if new_content.get("title"):
            existing["title"] = new_content["title"]
        if new_content.get("intro"):
            existing["intro"] = new_content["intro"]
        if new_content.get("sections"):
            # Replace or append sections
            existing["sections"] = new_content["sections"]
    else:
        content_data[content_key] = new_content
    
    with open(content_json_path, 'w', encoding='utf-8') as f:
        json.dump(content_data, f, indent=2, ensure_ascii=False)
    
    print(f"✓ Updated {content_json_path} with content for '{content_key}'")


def process_word_documents(word_docs_dir="word-docs", output_dir="."):
    """Process all Word documents in the specified directory."""
    word_docs_path = Path(word_docs_dir)
    
    if not word_docs_path.exists():
        print(f"Creating directory: {word_docs_dir}")
        word_docs_path.mkdir(parents=True, exist_ok=True)
        print(f"\nPlease place your .docx files in the '{word_docs_dir}' directory")
        print("Then run this script again.\n")
        return
    
    # Create assets directory for extracted images
    assets_dir = Path(output_dir) / "assets"
    assets_dir.mkdir(parents=True, exist_ok=True)
    
    content_json_path = Path(output_dir) / "assets" / "content.json"
    
    docx_files = list(word_docs_path.glob("*.docx"))
    
    if not docx_files:
        print(f"No .docx files found in '{word_docs_dir}' directory")
        return
    
    print(f"Found {len(docx_files)} Word document(s)\n")
    
    for docx_file in docx_files:
        print(f"Processing: {docx_file.name}")
        
        try:
            # Extract content
            content = extract_text_from_docx(docx_file)
            
            # Extract images
            images = extract_images_from_docx(docx_file, assets_dir)
            if images:
                print(f"  Extracted {len(images)} image(s)")
            
            # Determine content key from filename
            content_key = map_filename_to_content_key(docx_file.stem)
            
            if not content_key:
                print(f"  ⚠ Warning: Could not determine content key from filename '{docx_file.stem}'")
                print(f"  Available keys: maps, tectonics, weathering, fluvial, climate, about, home, references")
                content_key = input(f"  Enter content key for this document (or press Enter to skip): ").strip()
                if not content_key:
                    continue
            
            # Update content.json
            update_content_json(content_json_path, content, content_key)
            
            print(f"  ✓ Successfully processed '{docx_file.name}' -> '{content_key}'\n")
            
        except Exception as e:
            print(f"  ✗ Error processing '{docx_file.name}': {e}\n")
            import traceback
            traceback.print_exc()


if __name__ == "__main__":
    import sys
    
    # Allow custom directory paths
    word_docs_dir = sys.argv[1] if len(sys.argv) > 1 else "word-docs"
    output_dir = sys.argv[2] if len(sys.argv) > 2 else "."
    
    print("=" * 60)
    print("Word Document Content Extractor")
    print("=" * 60)
    print()
    
    process_word_documents(word_docs_dir, output_dir)
    
    print("=" * 60)
    print("Done!")
    print("=" * 60)

