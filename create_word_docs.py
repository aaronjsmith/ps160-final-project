#!/usr/bin/env python3
"""
Create Word documents from HTML files.

This script reads HTML files and converts them to Word documents (.docx)
with proper formatting, headings, and structure.
"""

import os
import re
from pathlib import Path
from html.parser import HTMLParser
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


class HTMLContentExtractor(HTMLParser):
    """Extract text content and structure from HTML."""
    
    def __init__(self):
        super().__init__()
        self.content = []
        self.current_tag = None
        self.current_text = []
        self.in_script = False
        self.in_style = False
        
    def handle_starttag(self, tag, attrs):
        if tag in ['script', 'style']:
            if tag == 'script':
                self.in_script = True
            elif tag == 'style':
                self.in_style = True
            return
        
        if tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'li', 'blockquote']:
            # Save any accumulated text before starting new element
            if self.current_text:
                text = ' '.join(self.current_text).strip()
                if text:
                    self.content.append(('text', text))
                self.current_text = []
            
            self.current_tag = tag
        elif tag == 'strong':
            self.current_text.append('**')
        elif tag == 'em':
            self.current_text.append('*')
    
    def handle_endtag(self, tag):
        if tag in ['script', 'style']:
            if tag == 'script':
                self.in_script = False
            elif tag == 'style':
                self.in_style = False
            return
        
        if tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'li', 'blockquote']:
            text = ' '.join(self.current_text).strip()
            if text:
                self.content.append((tag, text))
            self.current_text = []
            self.current_tag = None
        elif tag == 'strong':
            self.current_text.append('**')
        elif tag == 'em':
            self.current_text.append('*')
    
    def handle_data(self, data):
        if not self.in_script and not self.in_style:
            # Clean up whitespace
            cleaned = ' '.join(data.split())
            if cleaned:
                self.current_text.append(cleaned)


def extract_content_from_html(html_path):
    """Extract structured content from HTML file."""
    with open(html_path, 'r', encoding='utf-8') as f:
        html_content = f.read()
    
    # Remove HTML comments
    html_content = re.sub(r'<!--.*?-->', '', html_content, flags=re.DOTALL)
    
    # Extract content from main/article tag
    main_match = re.search(r'<main[^>]*>(.*?)</main>', html_content, re.DOTALL | re.IGNORECASE)
    article_match = re.search(r'<article[^>]*>(.*?)</article>', html_content, re.DOTALL | re.IGNORECASE)
    
    if article_match:
        content_html = article_match.group(1)
    elif main_match:
        content_html = main_match.group(1)
    else:
        # Fallback: extract from body
        body_match = re.search(r'<body[^>]*>(.*?)</body>', html_content, re.DOTALL | re.IGNORECASE)
        if body_match:
            content_html = body_match.group(1)
        else:
            content_html = html_content
    
    # Parse HTML
    parser = HTMLContentExtractor()
    parser.feed(content_html)
    
    return parser.content


def clean_text(text):
    """Clean extracted text."""
    # Remove markdown-style formatting
    text = text.replace('**', '')
    text = text.replace('*', '')
    # Clean up multiple spaces
    text = ' '.join(text.split())
    return text.strip()


def create_word_document(content, output_path):
    """Create a Word document from extracted content."""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    for tag, text in content:
        text = clean_text(text)
        if not text:
            continue
        
        # Skip navigation, footer, and other non-content elements
        if any(skip in text.lower() for skip in ['home', 'about', 'references', 'earth processes', 
                                                 'copyright', 'arizona study', '©']):
            if len(text) < 50:  # Likely navigation
                continue
        
        if tag == 'h1':
            para = doc.add_heading(text, level=1)
        elif tag == 'h2':
            para = doc.add_heading(text, level=2)
        elif tag == 'h3':
            para = doc.add_heading(text, level=3)
        elif tag == 'h4':
            para = doc.add_heading(text, level=4)
        elif tag == 'h5':
            para = doc.add_heading(text, level=5)
        elif tag == 'h6':
            para = doc.add_heading(text, level=6)
        elif tag == 'blockquote':
            para = doc.add_paragraph(text)
            para.style = 'Quote'
        elif tag == 'li':
            # Add as bullet point
            para = doc.add_paragraph(text, style='List Bullet')
        elif tag == 'p':
            para = doc.add_paragraph(text)
        else:
            # Default to paragraph
            para = doc.add_paragraph(text)
    
    doc.save(output_path)
    print(f"[OK] Created: {output_path}")


def process_html_files(html_dir=".", word_docs_dir="word-docs"):
    """Process all HTML files and create Word documents."""
    html_path = Path(html_dir)
    word_docs_path = Path(word_docs_dir)
    
    # Create word-docs directory
    word_docs_path.mkdir(parents=True, exist_ok=True)
    
    # HTML files to process (excluding index.html which is the home page)
    html_files = [
        ('maps-location-cartographers.html', 'maps-location-cartographers.docx'),
        ('plate-tectonics-earthquakes-volcanoes.html', 'plate-tectonics-earthquakes-volcanoes.docx'),
        ('weathering-mass-wasting-erosion.html', 'weathering-mass-wasting-erosion.docx'),
        ('fluvial-processes-oceans-coastlines.html', 'fluvial-processes-oceans-coastlines.docx'),
        ('climate-controls-biomes-climate-change.html', 'climate-controls-biomes-climate-change.docx'),
        ('about.html', 'about.docx'),
    ]
    
    print("=" * 60)
    print("Creating Word Documents from HTML Files")
    print("=" * 60)
    print()
    
    for html_file, docx_file in html_files:
        html_path_full = html_path / html_file
        docx_path_full = word_docs_path / docx_file
        
        if not html_path_full.exists():
            print(f"[SKIP] Skipping: {html_file} (not found)")
            continue
        
        print(f"\nProcessing: {html_file}")
        try:
            content = extract_content_from_html(html_path_full)
            create_word_document(content, docx_path_full)
        except Exception as e:
            print(f"[ERROR] Error processing {html_file}: {e}")
            import traceback
            traceback.print_exc()
    
    print()
    print("=" * 60)
    print("Done! Word documents created in 'word-docs' directory")
    print("=" * 60)


if __name__ == "__main__":
    import sys
    
    html_dir = sys.argv[1] if len(sys.argv) > 1 else "."
    word_docs_dir = sys.argv[2] if len(sys.argv) > 2 else "word-docs"
    
    process_html_files(html_dir, word_docs_dir)
